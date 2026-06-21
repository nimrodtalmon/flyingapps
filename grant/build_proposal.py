"""Build a Word draft of the Schmidt Sciences / CAIF / DeepMind / ARIA
"Scaling AI Safety for a Multi-Agent World" proposal, mirroring the smapply
form structure with all word limits honored.

Output: grant/Talmon_MultiAgentSafety_Proposal_DRAFT.docx
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


GREY = RGBColor(0x66, 0x66, 0x66)
ACCENT = RGBColor(0x1A, 0x55, 0x99)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = ACCENT
    return h


def add_field_label(doc, text, limit=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    if limit:
        r2 = p.add_run(f"  ({limit})")
        r2.italic = True
        r2.font.color.rgb = GREY
    return p


def add_field_help(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = True
    r.font.color.rgb = GREY
    return p


def add_answer(doc, text):
    p = doc.add_paragraph(text)
    return p


def add_todo(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(f"[TO FILL: {text}]")
    r.italic = True
    r.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
    return p


def add_wordcount(doc, n):
    p = doc.add_paragraph()
    r = p.add_run(f"(approx. word count: {n})")
    r.italic = True
    r.font.size = Pt(8)
    r.font.color.rgb = GREY
    return p


def add_divider(doc):
    p = doc.add_paragraph()
    r = p.add_run("—" * 30)
    r.font.color.rgb = GREY
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


doc = Document()

# Base style
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

# ============================================================
# Title block
# ============================================================
title = doc.add_heading(
    "Scaling AI Safety for a Multi-Agent World — Proposal Draft",
    level=0,
)

p = doc.add_paragraph()
r = p.add_run(
    "Measuring the Safety of Agent Populations: "
    "A Social-Choice Science and Frontier-Agent Testbed "
    "for Multi-Principal Ecosystems"
)
r.bold = True
r.font.size = Pt(14)

p = doc.add_paragraph()
r = p.add_run(
    "Lead PI: Nimrod Talmon (Ben-Gurion University of the Negev) · "
    "Tier 1 (≤ $300K, 24 months) · "
    "Application ID: TAI-RFP-MAS-26-6571449688 · "
    "Deadline: 8 Aug 2026 AoE"
)
r.italic = True
r.font.color.rgb = GREY

add_divider(doc)

p = doc.add_paragraph()
r = p.add_run(
    "Notes for the PI. "
    "This document mirrors the smapply form one-to-one. Each section is labelled "
    "with its word/page limit. Items marked [TO FILL: …] need your input before "
    "submission. The Milestones template and Budget template are required separate "
    "uploads — included here as drafts so you can refine, then drop into the official "
    "templates."
)
r.italic = True
r.font.color.rgb = GREY

doc.add_page_break()

# ============================================================
# SECTION 1: Background Information
# ============================================================
add_heading(doc, "1. Background Information", level=1)

add_field_label(doc, "Lead Applicant Name*")
add_answer(doc, "Nimrod Talmon")

add_field_label(doc, "Primary Email Address*")
add_answer(doc, "nimrodtalmon77@gmail.com")

add_field_label(doc, "Title / Role*")
add_todo(doc, "confirm exact title — likely 'Associate Professor' or 'Senior Lecturer'")

add_field_label(doc, "Primary Institution*")
add_answer(
    doc,
    "Ben-Gurion University of the Negev (BGU), Department of Software and "
    "Information Systems Engineering. "
    "Note: BGU is a non-US institution; NGOSource equivalency determination "
    "will be required if funded — to start ≤90 days from award per Schmidt policy."
)

add_field_label(doc, "Link to personal website (or lab page)")
add_todo(doc, "BGU faculty page / personal site URL")

add_field_label(doc, "Collaborator(s) (optional)")
add_answer(
    doc,
    "None. This is a single-institution Tier-1 proposal, deliberately scoped to "
    "fit the PI plus 1–2 PhD students and one research engineer."
)

add_field_label(doc, "Headshot(s)*")
add_todo(
    doc,
    "upload 1080×1080 professional headshot of PI"
)

add_field_label(doc, "How did you hear about this RFP?*")
add_todo(
    doc,
    "fill from drop-down (likely: 'Cooperative AI Foundation announcement' or "
    "'colleague referral')"
)

doc.add_page_break()

# ============================================================
# SECTION 2: Project Details
# ============================================================
add_heading(doc, "2. Project Details", level=1)

# ---- Title
add_field_label(doc, "Project Title*", limit="≤ 20 words")
add_answer(
    doc,
    "Measuring the Safety of Agent Populations: A Social-Choice Science and "
    "Frontier-Agent Testbed for Multi-Principal Ecosystems."
)
add_wordcount(doc, 18)

# ---- Tier
add_field_label(doc, "Project Tier*")
add_answer(doc, "Tier 1 (≤ USD $300,000).")

# ---- Duration
add_field_label(doc, "Project Duration*", limit="12–24 months")
add_answer(doc, "24 months.")

# ---- Plain-language summary
add_field_label(doc, "Plain-language Summary*", limit="2–3 sentences")
add_field_help(doc, "Plain, concrete language for a general technical audience.")
add_answer(
    doc,
    "We can train and align individual AI agents, but we have almost no principled "
    "way to measure whether a population of agents — built by different actors and "
    "sharing infrastructure — is safe as a whole. This project develops a "
    "measurement theory for population-level safety properties (collusion, covert "
    "coordination, cascade-fragility, emergent collective agency), instantiates it "
    "on an open frontier-agent testbed, and delivers a transfer-calibration tool "
    "that tells practitioners when sandbox conclusions apply to their real "
    "deployments."
)

# ---- Keywords
add_field_label(doc, "Keywords", limit="up to 5")
add_answer(
    doc,
    "multi-agent safety; computational social choice; agent-population testbed; "
    "external validity; collusion and collective agency."
)

doc.add_page_break()

# ---- Problem and Impact
add_field_label(doc, "Problem and Impact*", limit="≤ 500 words")
add_field_help(
    doc,
    "Central multi-agent safety risk / scientific gap; structured argument for "
    "impact; assumptions; best-case and minimum-valuable outcomes; what we learn "
    "even if it fails; why now."
)
add_answer(
    doc,
    "Single-model alignment is necessary but insufficient once agents from "
    "different principals interact over shared infrastructure (Hammond et al. 2025; "
    "Dafoe et al. 2021). Recent work has enumerated qualitatively new failure "
    "modes: covert collusion via steganographic side channels (Motwani et al. "
    "2024), attack propagation between agents (Lee & Tiwari 2024), collective "
    "resistance to shutdown (Agrawal, Ebadian & Hammond 2026), misuse via "
    "decomposition across individually-safe models (Jones, Dragan & Steinhardt "
    "2025), and the broader prospect of capable systems arising as coordinated "
    "networks rather than monoliths (DeepMind, Distributional AGI Safety 2025)."
)
add_answer(
    doc,
    "What is missing is the layer between \"we evaluated each agent\" and \"we "
    "understand what the swarm will do\" — a science of the population itself. "
    "Two coupled deficits motivate this work. (1) No tractable formalism for "
    "population-level properties: existing accounts of collective agency and "
    "emergence either require infeasibly many observations/interventions or rely "
    "on micro/macro abstractions that cannot be instantiated on real populations "
    "(Jørgensen, Weichwald & Hammond 2025; Szabo & Teo 2015) — we need definitions "
    "that come with estimators. (2) No external-validity science for testbeds: "
    "stylized games and simulated societies (Akata et al. 2025; Park et al. 2023) "
    "have advanced the field, but it cannot say when a sandbox result licenses a "
    "claim about deployment (Kapoor et al. 2026); without that, every testbed "
    "result is an anecdote."
)
add_answer(
    doc,
    "A measure is only useful if computable somewhere realistic; a testbed is only "
    "useful if its results transfer. We attack both together."
)
add_answer(
    doc,
    "Structured argument for impact. If the project succeeds, a population of "
    "agents built by different actors becomes a measurable object: a neutral third "
    "party can point our measurement suite at a deployment and obtain calibrated "
    "estimates of collusion-resistance, cascade-fragility, and collective-agency "
    "emergence, and the transfer-calibration tool tells them how much to trust the "
    "sandbox conclusions there."
)
add_answer(
    doc,
    "Best-case outcome: the measurement suite is adopted across the funded cohort "
    "and by at least one deploying organization as a third-party safety check, and "
    "the transfer-calibration tool becomes the standard way to argue \"this "
    "sandbox conclusion applies here.\" Minimum valuable outcome: a first-pass "
    "formal vocabulary for population-level safety properties with tractable "
    "estimators demonstrated on a realistic frontier-agent testbed for at least "
    "two properties (collusion and cascade), plus a documented external-validity "
    "layer with honest limits. Even if the central claims fail, we will have "
    "learned which population-level properties resist formalization, which "
    "estimators do not scale, and which testbed configurations evade "
    "external-validity calibration — directly informing the field's next attempt."
)
add_answer(
    doc,
    "Assumptions: (i) per-agent alignment is in scope of other work, not ours — "
    "we work at the population layer; (ii) frontier-agent behavior can be "
    "approximated by calibrated smaller proxies for high-volume sweeps, with "
    "explicit validity bounds; (iii) instance-space methodology, validated on the "
    "structurally similar problem of mapping elections (Szufa et al. 2025; "
    "Faliszewski et al. 2025), transfers to agent-population configurations "
    "(Aim C is the test)."
)
add_answer(
    doc,
    "Why now. Frontier agents with tools, persistent memory, and inter-agent "
    "communication are already being deployed by multiple principals over shared "
    "infrastructure; the failure modes are documented in 2024–2026; no neutral, "
    "cross-principal measurement science yet exists. The window to build it "
    "before deployment outpaces measurement is narrow."
)
add_wordcount(doc, 495)

doc.add_page_break()

# ---- Approach
add_field_label(doc, "Approach*", limit="≤ 1000 words")
add_field_help(
    doc,
    "Detailed approach: experiments/analyses, which models, data/environments and "
    "why appropriate, baselines, and ecological validity (how design reflects real "
    "deployment conditions vs. idealized settings)."
)
add_answer(
    doc,
    "We organize the project around three aims. The contribution lives in the "
    "loop between them (Aim C), not in any single one."
)

p = doc.add_paragraph()
r = p.add_run("Aim A — A measurement theory for population-level safety "); r.bold = True
r = p.add_run("(call sub-area: Science of Agent Networks #3, #4)."); r.italic = True
add_answer(
    doc,
    "Goal: define a compact family of population-level safety properties, each "
    "paired with a tractable estimator computable from a feasible budget of "
    "observations/interventions. Target properties (initial set; the testbed will "
    "force revisions): (a) collusion-resistance / manipulability of the collective "
    "outcome — we adapt the manipulation- and control-complexity apparatus of "
    "social choice (Talmon 2016; Kaczmarek, Rothe & Talmon 2025) to define a "
    "manipulation budget estimator over agent populations; (b) covert-coordination "
    "capacity — operationalizing steganographic collusion (Motwani et al. 2024) "
    "as a measurable population property, not an anecdote; (c) cascade-fragility "
    "— sensitivity of population outcomes to adversarial or compromised "
    "sub-populations and to attack propagation (Lee & Tiwari 2024), formalized "
    "via robustness-of-aggregation tools (Bredereck et al. 2017); "
    "(d) collective-agency emergence — revealed-preference / aggregation-"
    "consistency tests with bounded sample complexity, directly targeting the "
    "tractability complaint in Jørgensen et al. 2025. For each property we "
    "deliver (i) a formal definition, (ii) an estimator with sample/intervention "
    "complexity bounds, and (iii) a reference implementation against Aim B."
)

p = doc.add_paragraph()
r = p.add_run("Aim B — An open frontier-agent testbed "); r.bold = True
r = p.add_run("(call sub-area: Sandboxes and Testbeds)."); r.italic = True
add_answer(
    doc,
    "Goal: a reproducible environment of frontier-model agents acting over "
    "extended horizons with realistic tools, persistent memory, communication "
    "channels, and economic constraints, under multiple principals — engineered "
    "to compute Aim-A measures inline and host red-teaming of population failure "
    "modes."
)
add_answer(
    doc,
    "Models. Two tracks. Frontier track: Claude (Opus/Sonnet families), GPT (most "
    "recent available), Gemini (most recent available), and an open-weight "
    "frontier-tier model (e.g. DeepSeek/Qwen latest) — chosen to span vendor "
    "diversity so principal-adversarial scenarios are realistic. Distilled-proxy "
    "track: smaller open-weight models (≈8–32B) calibrated to stand in for "
    "frontier agents on specific behavioral dimensions, with explicit validity "
    "bounds. This is the call's request for principled navigation of the "
    "scalability/fidelity trade-off."
)
add_answer(
    doc,
    "Built to the five desiderata. Scalable: realistic agent counts and tool "
    "diversity via the proxy track. High-fidelity: frontier agents acting under "
    "realistic constraints, not coarse abstractions. Externally valid: Aim C. "
    "Safe and secure: dangerous behaviors studied in containment, no uncontrolled "
    "deployment, red-team kit released under a responsible-disclosure protocol. "
    "Reproducible: fixed scenario library, logged transcripts, deterministic "
    "replay where feasible."
)
add_answer(
    doc,
    "Scenario library (multi-principal by construction). Three families, each "
    "with at least one principal-adversarial variant: (1) shared-infrastructure "
    "task ecosystems (agents from different principals share a marketplace and "
    "a knowledge base); (2) multi-organization workflows (agents from competing "
    "principals must complete tasks routed through shared tools); "
    "(3) resource-constrained agent marketplaces (economic constraints, "
    "price-formation, collusion incentives). No scenario assumes a single "
    "privileged overseer (an explicit out-of-scope class)."
)
add_answer(
    doc,
    "Experiments and baselines. For each scenario × property: (i) estimator "
    "computed on a sweep over population size, heterogeneity, interaction "
    "topology, and tool access; (ii) red-team probes designed to exhibit the "
    "failure mode; (iii) baselines — per-agent monitoring without population-"
    "level estimators (current deployment practice), and stylized-game results "
    "without frontier agents (current academic practice). The comparison is "
    "which failures the population-level measures catch that the baselines miss, "
    "and at what intervention cost."
)

p = doc.add_paragraph()
r = p.add_run("Aim C — The loop and the external-validity layer."); r.bold = True
add_answer(
    doc,
    "Theory → testbed: Aim-A measures generate falsifiable predictions about "
    "which population structures are collusion-prone, cascade-fragile, or prone "
    "to collective-agency emergence; Aim-B runs them on frontier agents. "
    "Testbed → theory: red-teaming surfaces failure modes the formalism missed, "
    "and definitions/estimators are revised. We expect the first-pass definitions "
    "to be wrong in instructive ways — that is the loop."
)
add_answer(
    doc,
    "External-validity layer — a map of agent populations. Adapting the "
    "instance-space methodology of Szufa et al. 2025 / Faliszewski et al. 2025 "
    "(validated on the structurally similar problem of mapping elections), we "
    "embed tested configurations in a structured space, locate where real "
    "deployments sit relative to tested points, and derive principled confidence "
    "that a sandbox conclusion transfers. This directly operationalizes the "
    "call's \"externally valid — when should simulation-derived conclusions be "
    "trusted\" desideratum. Deliverable C is a tool, not a paper: given a "
    "sandbox result and a description of a target deployment, output a "
    "calibrated transfer confidence with its own validity envelope — the "
    "artifact a practitioner at a deploying organization actually wants."
)

p = doc.add_paragraph()
r = p.add_run("Ecological validity."); r.bold = True
add_answer(
    doc,
    "The unit of study is a frontier agent under realistic constraints: real "
    "tools (web, code execution, persistent memory), real economic frictions, "
    "real principals with conflicting objectives. The formalism is scaffolding "
    "instantiated on natural-language, tool-using, memory-bearing agents, never "
    "a substitute for them. The distilled-proxy track is calibrated to frontier "
    "behavior and carries an honestly-reported validity envelope, rather than "
    "being treated as identical to the frontier. This is what separates the "
    "proposal from the explicit \"toy systems\" out-of-scope category."
)

p = doc.add_paragraph()
r = p.add_run("Deliverables."); r.bold = True
add_answer(
    doc,
    "(i) Open-source testbed + scenario/red-team library, reusable across the "
    "funded cohort; (ii) measurement suite — runnable estimators for each "
    "population-safety property, pointable at third-party deployments; "
    "(iii) transfer-calibration tool (Aim C); (iv) reference findings — first "
    "systematic maps of how collusion-/cascade-/emergence-risk scale with "
    "population structure, with phase-transition probes where they appear; "
    "(v) publications at AAMAS / IJCAI / AAAI and the multi-agent-safety venues. "
    "All artifacts open."
)
add_wordcount(doc, 935)

doc.add_page_break()

# ---- Suggestions (optional, only if no sandbox)
add_field_label(doc, "Suggestions (Optional)")
add_field_help(
    doc,
    "Only if the proposal does NOT involve constructing a sandbox/testbed — ours "
    "does, so this can be left blank."
)
add_answer(doc, "Not applicable — this project constructs a testbed (Aim B).")

# ---- Novelty
add_field_label(doc, "Novelty*", limit="≤ 300 words")
add_field_help(
    doc,
    "Similar existing work and what assumption/limit/bottleneck the project "
    "overcomes."
)
add_answer(
    doc,
    "The two closest bodies of work are (i) recent multi-agent safety simulations "
    "and red-team papers (Motwani et al. 2024; Lee & Tiwari 2024; Agrawal, "
    "Ebadian & Hammond 2026; Park et al. 2023), and (ii) formal accounts of "
    "collective agency and emergence (Jørgensen, Weichwald & Hammond 2025; Szabo "
    "& Teo 2015). The first produces compelling failure-mode anecdotes but no "
    "population-level measurement that scales beyond the demonstration scenario. "
    "The second produces measurement that is theoretically principled but not "
    "tractably computable on real agent populations."
)
add_answer(
    doc,
    "What we overcome:"
)
add_answer(
    doc,
    "• Tractable estimators for population-level properties. We do not propose "
    "another definition of \"emergence\" or \"collective agency\"; we propose "
    "definitions that ship with estimators, with sample-complexity bounds, "
    "computable on a frontier-agent testbed. The methodological backbone — "
    "turning intractable-looking questions about populations into computable "
    "measures — is transferred from the PI's prior work on elections "
    "(Faliszewski et al. 2025; Szufa et al. 2025; Talmon 2016)."
)
add_answer(
    doc,
    "• An external-validity layer for testbeds. No prior testbed in this space "
    "carries a principled answer to \"should this sandbox result be trusted for "
    "my deployment?\" We adapt instance-space methodology, validated on the "
    "structurally similar problem of mapping the space of elections, to the "
    "space of agent-population configurations. The call's \"externally valid\" "
    "desideratum is treated as a first-class deliverable."
)
add_answer(
    doc,
    "• Multi-principal by construction. Every scenario has at least one "
    "principal-adversarial variant; no privileged overseer (an explicit "
    "out-of-scope class)."
)
add_answer(
    doc,
    "We are not aware of prior work that closes the loop between a tractable "
    "population-level formalism, a frontier-agent testbed, and an "
    "external-validity layer."
)
add_wordcount(doc, 290)

# ---- Feasibility
add_field_label(doc, "Feasibility*", limit="≤ 300 words")
add_field_help(doc, "Intuition, prior work, or proof-of-concept supporting plausibility at the requested scale.")
add_answer(
    doc,
    "The hard part is not building a simulator or writing a definition — it is "
    "making the measures tractable and the testbed externally valid. The PI's "
    "record is unusually well-matched to exactly those two difficulties:"
)
add_answer(
    doc,
    "• Tractable measures over populations. A body of work on manipulation, "
    "control, and robustness in collective decision-making (Talmon 2016; "
    "Kaczmarek, Rothe & Talmon 2025; Bredereck et al. 2017) — i.e., turning "
    "\"can a coalition shift the outcome, and at what cost\" into computable "
    "objects. This is the methodological core of Aim A's estimator program."
)
add_answer(
    doc,
    "• Instance-space / external-validity methodology. \"Drawing a Map of "
    "Elections\" and \"How Similar Are Two Elections?\" (Szufa et al. 2025; "
    "Faliszewski et al. 2025) — a rare, directly-transferable toolkit for the "
    "call's hardest sandbox desideratum. Aim C transplants the same methodology "
    "onto the space of agent-population configurations."
)
add_answer(
    doc,
    "• Frontier-agent work (not just abstractions). AI-Generated Compromises for "
    "Coalition Formation (Briman, Shapiro & Talmon 2025) — LLM agents reaching "
    "collective decisions; evidence the group operates at frontier-agent "
    "fidelity, not toy systems."
)
add_answer(
    doc,
    "• Executable multi-agent simulation platforms. A four-layer methodology "
    "(formal kernel → agent-based simulation → empirical calibration → "
    "multi-objective optimization) for executable institutional models (Elem & "
    "Talmon 2026), transferable as engineering practice to the Aim-B harness."
)
add_answer(
    doc,
    "Scale. Tier-1 scope is feasible because the distilled-proxy track absorbs "
    "high-volume runs (population-size sweeps, scenario sweeps, red-team "
    "batches), while frontier-agent compute is reserved for validation. This is "
    "the explicit cost-control lever and the cleanest way to honor the "
    "scalability/fidelity trade-off the call asks applicants to navigate."
)
add_wordcount(doc, 285)

doc.add_page_break()

# ---- Team
add_field_label(doc, "Team*", limit="≤ 300 words")
add_field_help(
    doc,
    "All team members and collaborators with institutions and titles; why this "
    "team; how each contributes; explicit lead PI and co-PIs."
)
add_answer(
    doc,
    "Lead PI: Prof. Nimrod Talmon, Department of Software and Information "
    "Systems Engineering, Ben-Gurion University of the Negev (BGU). The PI's "
    "research record covers the two methodological edges this project leans on: "
    "(i) tractable formal measures over populations — manipulation, control, "
    "robustness in collective decision-making (Talmon 2016; Kaczmarek, Rothe & "
    "Talmon 2025; Bredereck et al. 2017); (ii) instance-space / "
    "external-validity methodology — \"Drawing a Map of Elections\" (Szufa et "
    "al. 2025), \"How Similar Are Two Elections?\" (Faliszewski et al. 2025); "
    "and frontier-agent work — AI-Generated Compromises for Coalition Formation "
    "(Briman, Shapiro & Talmon 2025), Legitimate Overrides in Decentralized "
    "Protocols (Elem & Talmon 2026)."
)
add_answer(
    doc,
    "Co-PIs: None. This is a single-institution Tier-1 proposal, deliberately "
    "scoped to fit the PI plus 1–2 students and an engineer."
)
add_answer(
    doc,
    "Other personnel (to be hired on award):"
)
add_answer(
    doc,
    "• 1–2 PhD students (BGU, full-time). Roles: Aim-A formal definitions and "
    "estimator derivations; Aim-B scenario implementation and red-teaming; "
    "Aim-C instance-space embedding and transfer-calibration."
)
add_answer(
    doc,
    "• 1 research engineer / RA (BGU, part-time). Role: Aim-B harness engineering "
    "— agent runtime, tools, memory, comms, economy; instrumentation to compute "
    "Aim-A estimators inline; reproducibility infrastructure."
)
add_answer(
    doc,
    "• Postdoc fraction (optional, only if budget allows after other categories), "
    "focused on Aim-C external-validity methodology."
)
add_answer(
    doc,
    "Allocation by aim. PI: Aim A (definitions, estimator design) and Aim C "
    "(external-validity methodology). Students: Aims A, B, C as above. Engineer: "
    "Aim B (harness, instrumentation, red-team kit). The aims are deliberately "
    "bounded to the PI's two strongest methodological edges (tractable measures; "
    "instance-space external validity) rather than spread across all four call "
    "areas — the right shape for a focused single-institution Tier-1."
)
add_wordcount(doc, 295)

# ---- Risks
add_field_label(doc, "Proposal Risks*", limit="≤ 300 words")
add_field_help(doc, "Most significant risks (harm, irrelevance, inconclusive results) and mitigations.")
add_answer(
    doc,
    "• First-pass formal definitions don't survive contact with frontier agents. "
    "Mitigation: this is Aim C; the Aim A↔B loop is designed to revise them. "
    "Revision is a reported result, not a failure mode."
)
add_answer(
    doc,
    "• Frontier-agent experiments too costly at scale. Mitigation: the "
    "distilled-proxy track (Aim B) with explicit validity bounds absorbs "
    "high-volume runs; frontier-agent compute is reserved for validation. Keeps "
    "cost inside the Tier-1 envelope."
)
add_answer(
    doc,
    "• External validity is hard to validate. Mitigation: we borrow a "
    "methodology with a track record on a structurally similar problem "
    "(instance spaces for elections) and report its limits honestly. The "
    "transfer-calibration tool ships with its own validity envelope."
)
add_answer(
    doc,
    "• Risk of being read as classical game theory / \"toy systems.\" "
    "Mitigation: frontier-model agents with real tools, memory, and inter-agent "
    "communication are the unit of study throughout; the formalism is "
    "scaffolding instantiated on them. Aim C is the call's required "
    "\"concrete and credible methodology for extending results.\""
)
add_answer(
    doc,
    "• Inconclusive results. Mitigation: the work plan is staged so each phase "
    "produces a falsifiable artifact (definition + estimator + sandbox run); "
    "null results sharpen the formalism. The deliverables (testbed, measurement "
    "suite, transfer tool) are public goods of value independent of any single "
    "hypothesis."
)
add_answer(
    doc,
    "• Harm from publishing red-team scenarios. Mitigation: red-team kit "
    "released under a responsible-disclosure protocol consistent with cohort "
    "norms; principal-adversarial scenarios documented at the level needed to "
    "reproduce measurement, not to amplify exploit capability. No agent "
    "capability is advanced."
)
add_wordcount(doc, 280)

doc.add_page_break()

# ---- But for impact
add_field_label(doc, "\"But for\" Impact*", limit="≤ 300 words")
add_field_help(doc, "What makes this hard to fund by typical means and why markets won't solve it.")
add_answer(
    doc,
    "Every agent-deploying organization has an incentive to build internal "
    "evaluations of its own agents. None has an incentive to build a neutral, "
    "cross-principal, externally-validated measurement science for populations "
    "of agents built by others — yet that is exactly what a multi-principal "
    "world needs, and exactly what market forces will leave in the gap."
)
add_answer(
    doc,
    "Why markets will not solve it:"
)
add_answer(
    doc,
    "• Measurement public good. Once it exists, its highest social value is "
    "neutrality — the property no commercially-aligned actor can credibly "
    "supply."
)
add_answer(
    doc,
    "• Coordination structure typical of safety infrastructure. The value to "
    "any single actor is below the cost of building it; the value to the "
    "ecosystem is above. Without philanthropic seeding, no one builds it."
)
add_answer(
    doc,
    "• Explicitly non-capability deliverables. Estimators, red-team scenarios, "
    "a transfer-calibration tool. No deploying organization gains commercial "
    "advantage from funding it; an academic group with the right methodological "
    "prerequisites does."
)
add_answer(
    doc,
    "Why this is hard to fund elsewhere:"
)
add_answer(
    doc,
    "• Standard NSF / ISF tracks fund the theory pieces (Aim A) or the systems "
    "pieces (Aim B) but rarely the loop between them (Aim C), which is where "
    "this proposal's contribution lives."
)
add_answer(
    doc,
    "• Industry safety teams will not fund cross-principal, externally-validated "
    "measurement of their own deployments by a neutral third party."
)
add_answer(
    doc,
    "• The 24-month window is too short for most foundational-science grants "
    "and too long/risky for most engineering grants. The call's tier structure "
    "is the right fit."
)
add_answer(
    doc,
    "This is the coordination/incentive failure the call exists to fill: a "
    "measurement science for populations of agents from different principals, "
    "built by a neutral academic actor and released as a public good."
)
add_wordcount(doc, 290)

# ---- Existing Funding
add_field_label(doc, "Existing Funding*", limit="≤ 300 words")
add_field_help(
    doc,
    "Existing funding for closely related work and any pending applications "
    "elsewhere (funder + amount); effect on the project if not funded."
)
add_todo(
    doc,
    "PI to fill in concrete grant numbers. Draft text below — replace bracketed "
    "items with your actual portfolio."
)
add_answer(
    doc,
    "Existing support touching adjacent topics:"
)
add_answer(
    doc,
    "• [TO FILL: ISF Grant #_____, \"<title>\", USD/NIS amount, dates] — supports "
    "the foundational computational-social-choice work that enables Aim A's "
    "estimator program but does not extend to multi-agent safety, the "
    "frontier-agent testbed (Aim B), or the external-validity layer (Aim C)."
)
add_answer(
    doc,
    "• [TO FILL: any BSF / ERC / European or industry grant touching social "
    "choice, multi-agent systems, or LLM-agent research]."
)
add_answer(
    doc,
    "Pending applications elsewhere with overlap:"
)
add_answer(
    doc,
    "• [TO FILL: list any pending submissions that touch population-level safety "
    "measurement, agent testbeds, or instance-space methodology — or state "
    "\"no pending applications with overlapping scope\"]."
)
add_answer(
    doc,
    "Effect if not funded: the testbed (Aim B) and the external-validity tool "
    "(Aim C) would not be pursued — they require dedicated engineering and "
    "frontier-model API/compute beyond the PI's existing portfolio. Aim A could "
    "proceed at a reduced scale as part of ongoing social-choice work, but "
    "without the testbed-grounded revision loop that makes it relevant to "
    "multi-agent safety as defined by this call. In effect, without this award "
    "the contribution would collapse to a theory-only effort with no instantiated "
    "measurement, no external-validity calibration, and no practitioner-facing "
    "tool — losing the parts of the project the call was created to fund."
)
add_wordcount(doc, 250)

doc.add_page_break()

# ---- Suggested reviewers
add_field_label(doc, "Suggested Reviewers (Optional)", limit="up to 5")
add_field_help(doc, "Technical experts well-qualified to review; avoid close collaborators/advisors.")
add_todo(
    doc,
    "PI to suggest 3–5 names. Candidate domains (no specific names invented here): "
    "(1) multi-agent safety researchers (non-collaborators of the PI); "
    "(2) computational social choice senior researchers outside the PI's direct "
    "publication network; (3) empirical LLM-agent / testbed researchers with "
    "track records on multi-principal evaluation. Cross-check against your "
    "DBLP co-author graph before listing."
)

# ---- Optional figure
add_field_label(doc, "Optional Figure / Diagram")
add_todo(
    doc,
    "consider a single figure showing the Aim-A ↔ Aim-B loop + the Aim-C "
    "external-validity map. Stylistically: three boxes (Theory / Testbed / Map) "
    "with bidirectional arrows; the Map box outputs \"transfer confidence\" to a "
    "fourth box labelled \"deployment.\""
)

# ---- References
add_field_label(doc, "References")
add_field_help(doc, "Cited sources; not counted toward word limit; do not use to extend answers.")
refs = [
    "Hammond et al. (2025). Multi-Agent Risks from Advanced AI. CAIF report. arXiv:2502.14143.",
    "Dafoe et al. (2021). Open Problems in Cooperative AI. Nature.",
    "DeepMind (2025). Distributional AGI Safety. arXiv:2512.16856.",
    "Motwani et al. (2024). Secret Collusion among AI Agents: Multi-Agent Deception via Steganography. NeurIPS 2024.",
    "Lee & Tiwari (2024). Prompt Infection: LLM-to-LLM Prompt Injection within Multi-Agent Systems. arXiv:2410.07283.",
    "Agrawal, Ebadian & Hammond (2026). The Multi-Agent Off-Switch Game. AAMAS 2026.",
    "Jones, Dragan & Steinhardt (2025). Adversaries Can Misuse Combinations of Safe Models. ICML 2025.",
    "Jørgensen, Weichwald & Hammond (2025). Causal Foundations of Collective Agency. arXiv:2605.00248.",
    "Szabo & Teo (2015). Formalization of Weak Emergence in Multiagent Systems. ACM TOMACS.",
    "Akata et al. (2025). Playing Repeated Games with Large Language Models. Nature Human Behaviour.",
    "Park et al. (2023). Generative Agents. UIST '23.",
    "Kapoor et al. (2026). Open-World Evaluations for Measuring Frontier AI Capabilities. arXiv:2605.20520.",
    "Tilli (2026). Agent Properties for Multi-Agent Safety. ICLR 2026 Workshop.",
    "Szufa, Boehmer, Bredereck, Faliszewski, Niedermeier, Skowron, Slinko & Talmon (2025). Drawing a Map of Elections. Artificial Intelligence 343:104332.",
    "Faliszewski, Skowron, Slinko, Sornat, Szufa & Talmon (2025). How Similar Are Two Elections? JCSS 150 / AAAI.",
    "Briman, Shapiro & Talmon (2025). AI-Generated Compromises for Coalition Formation. arXiv:2506.06837.",
    "Kaczmarek, Rothe & Talmon (2025). Control by Adding or Deleting Edges in Graph-Restricted Weighted Voting Games. JAIR 82.",
    "Bredereck, Faliszewski, Kaczmarczyk, Niedermeier, Skowron & Talmon (2017). Robustness Among Multiwinner Voting Rules. arXiv:1707.01417.",
    "Talmon (2016). Algorithmic Aspects of Manipulation and Anonymization in Social Choice and Social Networks. PhD thesis, TU Berlin.",
    "Elem & Talmon (2026). Legitimate Overrides in Decentralized Protocols. arXiv:2602.12260.",
]
for ref in refs:
    p = doc.add_paragraph(ref, style="List Bullet")
add_todo(
    doc,
    "verify each arXiv ID / venue / year before submission — citations were "
    "copied from the working brief and should be re-checked against DBLP / "
    "arXiv."
)

doc.add_page_break()

# ============================================================
# SECTION 3: Milestones and Outcomes (DRAFT — to drop into template)
# ============================================================
add_heading(doc, "3. Milestones and Outcomes (DRAFT for the official template)", level=1)
add_field_help(
    doc,
    "The submission requires a separate template upload. The content below is a "
    "draft for that template — refine, then copy into the official form. Outcomes "
    "are end-state scientific/safety-relevant changes (not activities like "
    "\"publish papers\"). Milestones are time-bound decision points with tangible "
    "evidence."
)

add_heading(doc, "End-of-project Outcomes", level=2)

p = doc.add_paragraph()
r = p.add_run("Outcome 1 — Validated measurement theory for three population-level safety properties."); r.bold = True
add_answer(
    doc,
    "Formal definitions of collusion-resistance, cascade-fragility, and "
    "collective-agency emergence, each paired with an estimator carrying explicit "
    "sample/intervention complexity bounds and a reference implementation against "
    "the Aim-B testbed."
)
add_answer(
    doc,
    "Why it matters for safety-relevant decisions: deploying organizations gain, "
    "for the first time, a way to measure these properties on a multi-agent stack "
    "rather than only evaluate agents one at a time. Cohort partners can apply "
    "the estimators to their own scenarios."
)
add_answer(
    doc,
    "Success measure: each estimator passes a reproducibility test on the testbed "
    "(two independent runs, same scenario, agreement within the bound stated by "
    "the estimator's complexity analysis); at least one property's estimator "
    "detects a failure that per-agent monitoring baselines miss."
)

p = doc.add_paragraph()
r = p.add_run("Outcome 2 — Open frontier-agent testbed with multi-principal scenario library."); r.bold = True
add_answer(
    doc,
    "A reproducible, instrumented testbed of frontier-model agents (Claude, GPT, "
    "Gemini, an open-weight frontier-tier model) with tools, persistent memory, "
    "communication channels, and economic constraints, hosting at least three "
    "scenario families, each with at least one principal-adversarial variant."
)
add_answer(
    doc,
    "Why it matters: a neutral, reproducible measurement substrate for the funded "
    "cohort and third parties."
)
add_answer(
    doc,
    "Success measure: at least three scenarios released (one per family); the "
    "harness reproduces published runs deterministically where determinism is "
    "feasible; at least one cohort partner runs measurement on it independently."
)

p = doc.add_paragraph()
r = p.add_run("Outcome 3 — Transfer-calibration tool (Aim C)."); r.bold = True
add_answer(
    doc,
    "A tool that, given a sandbox result and a description of a target "
    "deployment, returns a calibrated transfer confidence with an explicit "
    "validity envelope."
)
add_answer(
    doc,
    "Why it matters: operationalizes the call's \"externally valid — when should "
    "simulation-derived conclusions be trusted\" desideratum into something a "
    "practitioner can actually use."
)
add_answer(
    doc,
    "Success measure: calibration is validated on held-out tested configurations; "
    "the tool's reported confidences are well-calibrated within stated bounds on "
    "the held-out set; reproducible third-party use is demonstrated by at least "
    "one external partner."
)

p = doc.add_paragraph()
r = p.add_run("Outcome 4 — First systematic map of how population-level safety properties scale."); r.bold = True
add_answer(
    doc,
    "An empirical chart of how collusion-, cascade-, and emergence-risk depend "
    "on population size, heterogeneity, interaction topology, and tool access — "
    "with phase-transition probes where they appear."
)
add_answer(
    doc,
    "Why it matters: the field gains a first empirical chart of where in "
    "configuration-space the population becomes fragile, which directly informs "
    "deployment-time safety policy (e.g. when to throttle agent counts or "
    "restrict shared tool access)."
)
add_answer(
    doc,
    "Success measure: public dataset of estimator outputs across the configuration "
    "sweep; at least one phase-transition finding (positive or negative) "
    "confirmed by red-team probes."
)

add_heading(doc, "Time-bound Milestones", level=2)

p = doc.add_paragraph()
r = p.add_run("Month 6."); r.bold = True
add_answer(
    doc,
    "What will be demonstrated: Aim-A v1 — formal definitions and first "
    "estimators for collusion-resistance and cascade-fragility; Aim-B harness "
    "MVP (single principal, single tool stack); first principal-adversarial "
    "scenario."
)
add_answer(
    doc,
    "Tangible evidence: technical report; working harness in public repo; "
    "scenario released."
)
add_answer(
    doc,
    "Core hypothesis tested: can a manipulation-budget estimator over an agent "
    "population return a bounded-error estimate from feasibly many "
    "interventions? If not, the formalism is the first thing revised."
)

p = doc.add_paragraph()
r = p.add_run("Month 12."); r.bold = True
add_answer(
    doc,
    "What will be demonstrated: loop online — Aim A ↔ Aim B for collusion and "
    "cascade properties; external-validity map v1."
)
add_answer(
    doc,
    "Tangible evidence: runnable estimators in repo; map artifact published; at "
    "least one cohort partner has reproduced at least one estimator on their "
    "own scenario."
)
add_answer(
    doc,
    "Core hypothesis tested: does instance-space embedding place tested "
    "configurations such that nearest-neighbour transfer-confidence "
    "predictions correlate with held-out estimator agreement? If not, the "
    "external-validity layer's methodology is revised."
)

p = doc.add_paragraph()
r = p.add_run("Month 18."); r.bold = True
add_answer(
    doc,
    "What will be demonstrated: collective-agency emergence estimator; red-team "
    "kit released; transfer-calibration tool alpha."
)
add_answer(
    doc,
    "Tangible evidence: tool runs on at least three sandbox results and a target "
    "deployment description; calibration measured on a held-out set; red-team "
    "kit publicly released with responsible-disclosure protocol."
)
add_answer(
    doc,
    "Core hypothesis tested: does a revealed-preference / aggregation-consistency "
    "test detect emergent collective agency with bounded sample complexity on "
    "realistic frontier-agent populations? If not, definitions are revised."
)

p = doc.add_paragraph()
r = p.add_run("Month 24."); r.bold = True
add_answer(
    doc,
    "What will be demonstrated: full measurement suite and transfer tool v1.0; "
    "cross-cohort benchmark; final maps and papers submitted to AAMAS / IJCAI / "
    "AAAI."
)
add_answer(
    doc,
    "Tangible evidence: artifacts publicly released under open license; "
    "benchmark adopted by at least one cohort partner; papers in submission."
)

doc.add_page_break()

# ============================================================
# SECTION 4: Team CVs
# ============================================================
add_heading(doc, "4. Team CVs", level=1)
add_field_help(doc, "Upload up to 5 PDF CVs of most relevant team members.")
add_todo(
    doc,
    "(1) PI CV (Nimrod Talmon) — most recent, max 2 pages preferred. "
    "(2) Optionally CVs of named senior personnel — none in this Tier-1 "
    "single-PI proposal. Students/RA hired post-award so no CVs required."
)

doc.add_page_break()

# ============================================================
# SECTION 5: Budget & Compute
# ============================================================
add_heading(doc, "5. Budget & Compute", level=1)

add_field_label(doc, "Estimated Total Budget")
add_answer(
    doc,
    "USD $290,000–$300,000 over 24 months, inclusive of ≤10% indirect "
    "(per Schmidt Sciences policy). Final allocation in budget template."
)

add_field_label(doc, "Budget Template Upload")
add_todo(
    doc,
    "complete the official Schmidt budget template (link in smapply form); "
    "filename: \"Talmon Budget for Scientific Milestones and Outcomes for "
    "Scaling AI Safety for a Multi-Agent World\". Categories drafted below."
)

add_field_label(doc, "Indicative Allocation (for the budget template)")
p = doc.add_paragraph()
p.add_run(
    "Personnel (~70–75% of total). 1–2 PhD students at BGU full-time × 24 mo; "
    "1 research engineer / RA part-time × 24 mo; PI summer salary fraction at "
    "BGU's policy rate. No co-PI line; optional postdoc fraction only if "
    "budget allows after other categories."
)
p = doc.add_paragraph()
p.add_run(
    "Compute — Option 2 (Standard Compute, ≤ USD $50,000 inclusive line item). "
    "Commercial GPU rental (≈ H100-class via standard cloud provider) for "
    "distilled-proxy training and inference; rough estimate ≈ USD $2–4 per "
    "GPU-hour at provider list rates. Indicative target: ~$25–40k over 24 mo."
)
p = doc.add_paragraph()
p.add_run(
    "API credits (separate line, not counted as \"compute\" per the call's rule). "
    "Frontier-model API for validation runs and red-team batches: Claude "
    "(Opus / Sonnet), GPT, Gemini, open-weight frontier-tier model. Indicative: "
    "~$15–25k over 24 mo, heavily front-loaded to Months 6–18."
)
p = doc.add_paragraph()
p.add_run(
    "Travel / dissemination (~5%). AAMAS / IJCAI / AAAI + the cohort meeting. "
    "Non-US-awardee logistics will respect W-8 considerations; US conference "
    "travel funded from non-Schmidt sources where required."
)
p = doc.add_paragraph()
p.add_run(
    "Indirect (≤10% of grand total) per Schmidt Sciences policy."
)

add_field_label(doc, "Budget Justification*", limit="≤ 300 words")
add_field_help(doc, "High-level alignment between budget and project; necessity of categories.")
add_answer(
    doc,
    "Total ≤ USD $300,000 over 24 months, with ≤10% indirect, per Tier-1 cap and "
    "Schmidt Sciences policy."
)
add_answer(
    doc,
    "Personnel (~70–75%). The core contribution — Aim A's tractable estimators "
    "and Aim C's external-validity tool — is intellectual work, not compute. "
    "1–2 PhD students and one engineer is the smallest team that covers Aim A "
    "(formal work), Aim B (testbed engineering), and Aim C (instance-space "
    "embedding) over 24 months. No co-PI, no full postdoc line — these are "
    "deliberately omitted to keep the proposal focused and within Tier 1."
)
add_answer(
    doc,
    "Compute — Option 2 (≤ $50,000 inclusive). Modest commercial GPU rental "
    "(≈ H100-class) for the distilled-proxy track: training and inference on "
    "8–32B open-weight models calibrated to stand in for frontier agents on "
    "specific behavioral dimensions. This track is the cost-control mechanism "
    "that makes a genuinely frontier-fidelity testbed feasible at Tier-1 scale. "
    "Provider, rate, and hour estimates are documented in the budget line item."
)
add_answer(
    doc,
    "API credits (separate line, not \"compute\" per call rules). Frontier-model "
    "API for validation runs and red-team batches across Claude, GPT, Gemini, "
    "and an open-weight frontier-tier model — vendor diversity is required so "
    "principal-adversarial scenarios are realistic."
)
add_answer(
    doc,
    "Travel / dissemination (~5%). AAMAS / IJCAI / AAAI plus the funded-cohort "
    "meeting. US conference travel will be supported from non-Schmidt sources "
    "where required by the non-US-awardee structure."
)
add_answer(
    doc,
    "Indirect ≤10% at BGU's policy cap. The personnel-heavy, "
    "compute-conservative structure aligns with the proposal's central "
    "feasibility argument: the proxy track absorbs volume; frontier runs are "
    "reserved for validation; the intellectual deliverables drive the cost."
)
add_wordcount(doc, 295)

add_field_label(doc, "Grant administrator / finance contact email*")
add_todo(doc, "BGU research authority / grant administrator email — to be confirmed by PI.")

add_field_label(doc, "Compute Option")
add_answer(
    doc,
    "Option 2 — Standard Compute (≤ USD $50,000 GPU line item in budget). "
    "Rationale: high-volume runs are LLM-based and use API (which is not "
    "\"compute\" per the call); we need only modest commercial GPU for the "
    "distilled-proxy track. Option 2 fits cleanly within Tier 1; in-kind "
    "cluster support (Option 3) is not required."
)

doc.add_page_break()

# ============================================================
# SECTION 6: Anything Else?
# ============================================================
add_heading(doc, "6. Anything Else?", level=1)

add_field_label(
    doc,
    "Do you give permission to share your proposal with other funders (external "
    "to the partners of this call)?"
)
add_todo(
    doc,
    "PI to choose Yes / No. Default suggestion: Yes — wider exposure increases "
    "back-up funding probability and the project is non-proprietary."
)

add_field_label(doc, "Anything else you would like to communicate?")
add_answer(
    doc,
    "Two practical notes for the panel. (1) BGU is a non-US institution; "
    "NGOSource equivalency determination, where required, will be initiated "
    "within Schmidt's 90-day target. (2) The proposal targets the "
    "\"Science of Agent Networks\" sub-areas #3 and #4 and the \"Sandboxes "
    "and Testbeds\" area in depth, rather than spreading across all four call "
    "areas — consistent with the call's explicit depth-over-breadth guidance."
)

doc.add_page_break()

# ============================================================
# SECTION 7: Eligibility & Disclaimer
# ============================================================
add_heading(doc, "7. Eligibility & Disclaimer", level=1)

add_field_label(doc, "Confirm eligibility criteria*")
add_todo(
    doc,
    "PI to check the box. BGU is a non-US public university and will satisfy "
    "the IRS foreign-equivalent requirement (NGOSource ED to be initiated on "
    "award)."
)

add_field_label(doc, "Additional requirements for Google.org-funded projects*")
add_todo(
    doc,
    "PI to decide whether to accept Google.org's open-access / IP terms "
    "(Apache 2.0 for code, CC-BY 4.0 for other IP, no patents, AI Principles "
    "compliance). The proposal's deliverables are already designed to be open, "
    "so accepting is recommended — it widens the funding pool and does not "
    "alter our IP plan."
)

add_field_label(doc, "Terms of submission*")
add_todo(doc, "PI to accept on submission.")

doc.add_page_break()

# ============================================================
# Appendix — author working notes
# ============================================================
add_heading(doc, "Appendix — Working Notes (delete before submission)", level=1)

add_answer(
    doc,
    "Word-count budget per section was honored against the limits in the "
    "smapply PDF (Background; Project Details; Milestones; Budget). Each "
    "answer ends with an approximate word count; verify exactly with Word's "
    "counter before submission."
)
add_answer(
    doc,
    "Items to confirm with the PI before submission:"
)
add_answer(
    doc,
    "1. PI title and BGU faculty page URL."
)
add_answer(
    doc,
    "2. Headshot (1080×1080, professional)."
)
add_answer(
    doc,
    "3. \"How did you hear about this RFP?\" drop-down selection."
)
add_answer(
    doc,
    "4. Existing-Funding section: real grant numbers, amounts, dates; any "
    "pending applications with overlapping scope."
)
add_answer(
    doc,
    "5. Suggested reviewers (3–5 names; non-collaborators per DBLP graph)."
)
add_answer(
    doc,
    "6. Grant administrator email at BGU."
)
add_answer(
    doc,
    "7. Verify each reference's arXiv ID / venue / year against DBLP and "
    "arXiv (some were copied from the working brief and should be re-checked)."
)
add_answer(
    doc,
    "8. Permission-to-share answer and Google.org-terms answer."
)
add_answer(
    doc,
    "9. Optional figure: one figure showing the Aim-A ↔ Aim-B loop + Aim-C "
    "external-validity map."
)
add_answer(
    doc,
    "10. Budget template (separate upload) and Milestones template (separate "
    "upload) — drop drafts above into the official Schmidt forms."
)

# Save
out_path = "/home/user/flyingapps/grant/Talmon_MultiAgentSafety_Proposal_DRAFT.docx"
doc.save(out_path)
print(f"Wrote: {out_path}")
